"""
Document Loader Module
Handles loading and extracting text from PDF, TXT, and DOCX files.
"""

import os
from typing import List, Optional
from langchain_core.documents import Document


class DocumentLoader:
    """Handles document loading and text extraction from various file formats."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.docx'}
    
    def __init__(self):
        """Initialize the DocumentLoader."""
        pass
    
    def load_document(self, file_path: str) -> List[Document]:
        """
        Load a document and extract its text content.
        
        Args:
            file_path: Path to the document file.
            
        Returns:
            List of Document objects containing the extracted text.
            
        Raises:
            ValueError: If the file format is not supported.
            FileNotFoundError: If the file doesn't exist.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {ext}. Supported formats: {self.SUPPORTED_EXTENSIONS}")
        
        if ext == '.pdf':
            return self._load_pdf(file_path)
        elif ext == '.txt':
            return self._load_txt(file_path)
        elif ext == '.docx':
            return self._load_docx(file_path)
        
        return []
    
    def _load_pdf(self, file_path: str) -> List[Document]:
        """Load and extract text from a PDF file."""
        from pypdf import PdfReader
        
        documents = []
        reader = PdfReader(file_path)
        
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text.strip():
                doc = Document(
                    page_content=text,
                    metadata={
                        "source": file_path,
                        "page": page_num + 1,
                        "file_type": "pdf"
                    }
                )
                documents.append(doc)
        
        return documents
    
    def _load_txt(self, file_path: str) -> List[Document]:
        """Load and extract text from a TXT file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        if text.strip():
            return [Document(
                page_content=text,
                metadata={
                    "source": file_path,
                    "file_type": "txt"
                }
            )]
        return []
    
    def _load_docx(self, file_path: str) -> List[Document]:
        """Load and extract text from a DOCX file."""
        from docx import Document as DocxDocument
        
        doc = DocxDocument(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        full_text = "\n".join(text_parts)
        
        if full_text.strip():
            return [Document(
                page_content=full_text,
                metadata={
                    "source": file_path,
                    "file_type": "docx"
                }
            )]
        return []
    
    def load_multiple_documents(self, file_paths: List[str]) -> List[Document]:
        """
        Load multiple documents at once.
        
        Args:
            file_paths: List of paths to document files.
            
        Returns:
            Combined list of Document objects from all files.
        """
        all_documents = []
        
        for file_path in file_paths:
            try:
                docs = self.load_document(file_path)
                all_documents.extend(docs)
            except Exception as e:
                print(f"Error loading {file_path}: {str(e)}")
                continue
        
        return all_documents
    
    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """Check if a file format is supported."""
        ext = os.path.splitext(file_path)[1].lower()
        return ext in cls.SUPPORTED_EXTENSIONS
